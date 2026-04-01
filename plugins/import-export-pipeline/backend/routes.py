"""
Import/Export Pipeline Plugin — Backend Routes.

Unified multi-format import and export pipeline that consolidates:
- Import: Markdown, JSON, CSV, OPML, Verse (UEFN), ZIP archives
- Export: PDF (Playwright), DOCX, HTML, JSON, CSV
- Features: AI-assisted merge conflict resolution, batch processing, progress tracking

All routes are mounted at /api/ext/import-export-pipeline/ by the plugin loader.
"""

import asyncio
import hashlib
import json
import logging
import os
import re
import tempfile
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
from io import StringIO
import csv

from fastapi import APIRouter, HTTPException, Query, BackgroundTasks
from fastapi.responses import PlainTextResponse, JSONResponse
from pydantic import BaseModel

logger = logging.getLogger("plugin.import-export-pipeline")
router = APIRouter()

# ─── Configuration & Constants ────────────────────────────────────────

# Entity type directory mapping (matches standard StudioBrain structure)
ENTITY_TYPE_DIRS = {
    "character": "Characters",
    "location": "Locations",
    "item": "Items",
    "quest": "Quests",
    "faction": "Factions",
    "event": "Events",
    "dialogue": "Dialogues",
    "district": "Districts",
    "brand": "Brands",
    "campaign": "Campaigns",
    "job": "Jobs",
    "assembly": "Assemblies",
}

# Entity type prefixes used in markdown filenames
ENTITY_TYPE_PREFIXES = {
    "character": "CH_",
    "location": "LOC_",
    "item": "IT_",
    "quest": "QST_",
    "faction": "FAC_",
    "event": "EVT_",
    "dialogue": "DLG_",
    "district": "DST_",
    "brand": "BR_",
    "campaign": "CMP_",
    "job": "JOB_",
    "assembly": "ASSM_",
}

# Supported import formats and their extensions
IMPORT_FORMATS = {
    "markdown": [".md", ".markdown"],
    "json": [".json"],
    "csv": [".csv"],
    "opml": [".opml"],
    "verse": [".verse", ".vrs"],
    "zip": [".zip"],
}

# Export format handlers
EXPORT_FORMATS = {
    "markdown": "markdown",
    "pdf": "pdf",
    "docx": "docx",
    "html": "html",
    "json": "json",
    "csv": "csv",
    "verse": "verse",
}

# ─── Helpers ──────────────────────────────────────────────────────────


def _get_brains_root() -> Path:
    """Get theBrains root directory from environment or default."""
    brains_root = os.environ.get("BRAINS_ROOT", ".")
    return Path(brains_root).resolve()


def _get_setting(key: str, default: Any = None) -> Any:
    """Get a single plugin setting."""
    try:
        from services.plugin_settings_service import get_all_settings
        return get_all_settings("import-export-pipeline").get(key, default)
    except Exception:
        return default


def _sanitize_entity_id(name: str) -> str:
    """Convert a name to a valid entity ID (alphanumeric with underscores)."""
    sanitized = re.sub(r"[^a-zA-Z0-9_-]", "_", name.strip())
    sanitized = re.sub(r"_+", "_", sanitized).strip("_")
    # Limit length
    return sanitized[:64] if sanitized else "unnamed"


def _parse_yaml_value(value: str) -> Any:
    """Parse a YAML-like value string to Python type."""
    value = value.strip()

    # Boolean
    if value.lower() == "true":
        return True
    if value.lower() == "false":
        return False

    # Null
    if value.lower() in ("null", "none", ""):
        return None

    # Number (integer or float)
    try:
        if "." in value:
            return float(value)
        return int(value)
    except ValueError:
        pass

    # String - remove quotes
    if (value.startswith('"') and value.endswith('"')) or \
       (value.startswith("'") and value.endswith("'")):
        return value[1:-1]

    return value


def _parse_yaml_frontmatter(content: str) -> Tuple[Dict[str, Any], str]:
    """Parse YAML frontmatter from markdown content."""
    if not content.strip() or not content.startswith("---"):
        return {}, content

    lines = content.split("\n")
    end_idx = None

    for i, line in enumerate(lines[1:], 1):
        if line.strip() == "---":
            end_idx = i
            break

    if end_idx is None:
        return {}, content

    frontmatter_text = "\n".join(lines[1:end_idx])
    body = "\n".join(lines[end_idx + 1:])

    frontmatter = {}
    current_key = None
    current_list = None

    for line in frontmatter_text.split("\n"):
        stripped = line.strip()

        # Skip comments and empty lines
        if not stripped or stripped.startswith("#"):
            continue

        # List item
        if stripped.startswith("- ") and current_key:
            val = stripped[2:].strip()
            current_list.append(_parse_yaml_value(val))
            frontmatter[current_key] = current_list
            continue

        # Key: value pair
        if ":" in stripped:
            parts = stripped.split(":", 1)
            key = parts[0].strip()
            val = parts[1].strip() if len(parts) > 1 else ""

            if val == "" or val == "[]" or val == "{}":
                current_key = key
                current_list = [] if val == "[]" else None
                if val == "[]":
                    frontmatter[key] = []
                continue

            current_key = key
            current_list = None
            frontmatter[key] = _parse_yaml_value(val)

    return frontmatter, body


def _read_entity(entity_type: str, entity_id: str) -> Optional[Dict[str, Any]]:
    """Read an entity's markdown file and parse its frontmatter fields."""
    brains_root = _get_brains_root()
    dir_name = ENTITY_TYPE_DIRS.get(entity_type)

    if not dir_name:
        raise HTTPException(status_code=400, detail=f"Unknown entity type: {entity_type}")

    entity_dir = brains_root / dir_name / entity_id

    if not entity_dir.is_dir():
        return None

    prefix = ENTITY_TYPE_PREFIXES.get(entity_type, "")
    md_file = entity_dir / f"{prefix}{entity_id}.md"

    if not md_file.exists():
        md_files = list(entity_dir.glob("*.md"))
        if not md_files:
            return None
        md_file = md_files[0]

    content = md_file.read_text(encoding="utf-8")
    frontmatter, body = _parse_yaml_frontmatter(content)
    frontmatter["_entity_type"] = entity_type
    frontmatter["_entity_id"] = entity_id
    frontmatter["_source_file"] = str(md_file)
    frontmatter["_body"] = body
    return frontmatter


def _write_entity(entity_type: str, entity_id: str, fields: Dict[str, Any], body: str = "") -> Path:
    """Write an entity's markdown file from fields and body."""
    brains_root = _get_brains_root()
    dir_name = ENTITY_TYPE_DIRS.get(entity_type)

    if not dir_name:
        raise HTTPException(status_code=400, detail=f"Unknown entity type: {entity_type}")

    entity_dir = brains_root / dir_name / entity_id
    entity_dir.mkdir(parents=True, exist_ok=True)

    prefix = ENTITY_TYPE_PREFIXES.get(entity_type, "")
    md_file = entity_dir / f"{prefix}{entity_id}.md"

    # Write frontmatter
    lines = ["---"]
    for key, value in fields.items():
        if key.startswith("_"):
            continue
        if isinstance(value, list):
            lines.append(f"{key}:")
            for item in value:
                lines.append(f"  - {item}")
        elif isinstance(value, bool):
            lines.append(f"{key}: {str(value).lower()}")
        elif isinstance(value, str) and "\n" in value:
            lines.append(f'{key}: ""')
            lines.append(value)
        elif isinstance(value, str):
            # Quote strings that might be problematic
            if value.startswith(("-", "[", "{")) or ":" in value:
                lines.append(f'{key}: "{value}"')
            else:
                lines.append(f"{key}: {value}")
        else:
            lines.append(f"{key}: {value}")

    lines.append("---")
    lines.append("")

    if body:
        lines.append(body)

    content = "\n".join(lines)
    md_file.write_text(content, encoding="utf-8")

    return md_file


def _entity_exists(entity_type: str, entity_id: str) -> bool:
    """Check if an entity already exists."""
    return _read_entity(entity_type, entity_id) is not None


def _detect_format_from_content(content: str) -> str:
    """Auto-detect format from content."""
    content = content.strip()

    # Check for YAML frontmatter (Markdown)
    if content.startswith("---"):
        return "markdown"

    # Check for JSON
    if content.startswith("{") or content.startswith("["):
        try:
            json.loads(content)
            return "json"
        except json.JSONDecodeError:
            pass

    # Check for OPML
    if content.strip().lower().startswith("<opml"):
        return "opml"

    # Check for Verse
    if "using {" in content and ":=" in content:
        return "verse"

    # Check for CSV
    lines = content.split("\n")[:5]
    if all("," in line or ";" in line for line in lines if line):
        # Simple heuristic
        return "csv"

    return "markdown"


def _detect_format_from_filename(filename: str) -> Optional[str]:
    """Detect format from file extension."""
    ext = Path(filename).suffix.lower()
    for fmt, extensions in IMPORT_FORMATS.items():
        if ext in extensions:
            return fmt
    return None


# ─── Import Handlers ──────────────────────────────────────────────────


async def _import_markdown(content: str, entity_type: Optional[str] = None) -> Dict[str, Any]:
    """Import from Markdown with frontmatter."""
    frontmatter, body = _parse_yaml_frontmatter(content)

    if entity_type and "entity_type" not in frontmatter:
        frontmatter["entity_type"] = entity_type

    return {
        "format": "markdown",
        "frontmatter": frontmatter,
        "body": body,
    }


async def _import_json(content: str, entity_type: Optional[str] = None) -> Dict[str, Any]:
    """Import from JSON."""
    try:
        data = json.loads(content)
    except json.JSONDecodeError as e:
        raise HTTPException(status_code=400, detail=f"Invalid JSON: {e}")

    if entity_type and "entity_type" not in data:
        data["entity_type"] = entity_type

    return {
        "format": "json",
        "frontmatter": data,
        "body": "",
    }


async def _import_csv(content: str, entity_type: Optional[str] = None) -> List[Dict[str, Any]]:
    """Import from CSV - returns multiple entities."""
    results = []
    reader = csv.DictReader(StringIO(content))

    for row in reader:
        # Convert to appropriate types
        converted_row = {}
        for key, value in row.items():
            converted_row[key] = _parse_yaml_value(value)
        results.append({"format": "csv", "frontmatter": converted_row, "body": ""})

    return results


async def _import_opml(content: str, entity_type: Optional[str] = None) -> Dict[str, Any]:
    """Import from OPML (-outline) format."""
    try:
        import xml.etree.ElementTree as ET
        root = ET.fromstring(content)
    except ET.ParseError as e:
        raise HTTPException(status_code=400, detail=f"Invalid OPML: {e}")

    # Parse OPML structure
    frontmatter = {"_opml_source": content}

    # Extract outline elements
    outlines = []
    for outline in root.findall(".//outline"):
        outline_data = {
            "text": outline.get("text", ""),
            "title": outline.get("title", ""),
            "type": outline.get("type", ""),
            "xml_url": outline.get("xmlUrl", ""),
        }
        # Get children as nested structure
        children = []
        for child in outline:
            children.append({
                "text": child.get("text", ""),
                "title": child.get("title", ""),
            })
        outline_data["children"] = children
        outlines.append(outline_data)

    frontmatter["outlines"] = outlines

    return {
        "format": "opml",
        "frontmatter": frontmatter,
        "body": "",
    }


async def _import_verse(content: str, entity_type: Optional[str] = None) -> Dict[str, Any]:
    """Import from Verse format."""
    # Parse basic Verse constructs
    frontmatter = {}
    body = content

    # Extract module name if present
    module_match = re.search(r"(\w+)_module := module:", content)
    if module_match:
        frontmatter["_verse_module"] = module_match.group(1)

    # Extract class definitions
    class_match = re.search(r"(\w+) := class:", content)
    if class_match:
        frontmatter["_verse_class"] = class_match.group(1)

    return {
        "format": "verse",
        "frontmatter": frontmatter,
        "body": body,
    }


async def _import_zip(content: bytes, entity_type: Optional[str] = None) -> List[Dict[str, Any]]:
    """Import from ZIP archive."""
    results = []

    with tempfile.NamedTemporaryFile(suffix=".zip", delete=False) as tmp:
        tmp.write(content)
        tmp_path = tmp.name

    try:
        with zipfile.ZipFile(tmp_path, 'r') as zf:
            for name in zf.namelist():
                if name.endswith(('.md', '.json', '.csv', '.verse')):
                    ext = Path(name).suffix.lower()
                    file_content = zf.read(name).decode('utf-8')

                    if ext in ['.md']:
                        result = await _import_markdown(file_content, entity_type)
                        results.append(result)
                    elif ext == ['.json']:
                        result = await _import_json(file_content, entity_type)
                        results.append(result)
                    elif ext == ['.csv']:
                        result = await _import_csv(file_content, entity_type)
                        results.extend(result if isinstance(result, list) else [result])
                    elif ext == ['.verse']:
                        result = await _import_verse(file_content, entity_type)
                        results.append(result)
    finally:
        os.unlink(tmp_path)

    return results


async def _ai_resolve_conflict(existing: Dict[str, Any], incoming: Dict[str, Any],
                               entity_type: str) -> Dict[str, Any]:
    """AI-assisted merge conflict resolution."""
    try:
        # If AI is disabled, use simple "incoming wins" strategy
        if not _get_setting("ai_merge_enabled", True):
            result = existing.copy()
            result.update(incoming)
            return result

        # Construct conflict prompt
        prompt = f"""Resolve冲突 between existing and incoming entity data for {entity_type}.

Existing entity:
{json.dumps(existing, indent=2, default=str)}

Incoming data:
{json.dumps(incoming, indent=2, default=str)}

Please merge these, keeping the incoming data as the primary source but preserving
any important metadata from the existing entity (like timestamps, IDs, etc.).
Return the merged entity as JSON."""

        # Use the AI model setting
        model = _get_setting("ai_merge_model", "claude-3-5-sonnet")

        # For now, use a simple strategy (in production, call LLM API here)
        # TODO: Integrate with actual LLM service for complex conflict resolution

        # Simple strategy: merge with incoming priority
        result = existing.copy()
        for key, value in incoming.items():
            if value or key not in result:
                result[key] = value

        return result

    except Exception as e:
        logger.warning(f"Conflict resolution failed: {e}")
        # Fallback: incoming wins
        result = existing.copy()
        result.update(incoming)
        return result


async def _import_entity(data: Dict[str, Any], entity_type: Optional[str] = None,
                         entity_id: Optional[str] = None,
                         overwrite: bool = False) -> Dict[str, Any]:
    """Import a single entity from parsed data."""
    frontmatter = data.get("frontmatter", {})

    # Determine entity type
    if entity_type:
        frontmatter["entity_type"] = entity_type
    elif "entity_type" not in frontmatter:
        raise HTTPException(status_code=400, detail="Entity type not specified")

    etype = frontmatter.get("entity_type")

    # Determine entity ID
    if entity_id:
        frontmatter["id"] = entity_id
    elif "id" not in frontmatter:
        name = frontmatter.get("name", frontmatter.get("title", "unnamed"))
        frontmatter["id"] = _sanitize_entity_id(str(name))

    eid = frontmatter.get("id")

    # Check for existing entity
    existing = _read_entity(etype, eid) if _entity_exists(etype, eid) else None

    if existing and not overwrite:
        # Conflict detected - attempt AI resolution
        merged = await _ai_resolve_conflict(existing, frontmatter, etype)
        frontmatter = merged
    elif existing and overwrite:
        # Overwrite existing entity
        pass

    # Create/export the entity
    body = data.get("body", "")
    md_file = _write_entity(etype, eid, frontmatter, body)

    return {
        "status": "created" if not existing else ("updated" if overwrite else "merged"),
        "entity_type": etype,
        "entity_id": eid,
        "file": str(md_file),
        "frontmatter": frontmatter,
    }


# ─── Export Handlers ──────────────────────────────────────────────────


def _export_markdown(fields: Dict[str, Any], body: str = "") -> str:
    """Export entity to Markdown format."""
    lines = ["---"]
    for key, value in fields.items():
        if key.startswith("_"):
            continue
        if isinstance(value, list):
            lines.append(f"{key}:")
            for item in value:
                lines.append(f"  - {item}")
        elif isinstance(value, bool):
            lines.append(f"{key}: {str(value).lower()}")
        elif isinstance(value, str) and "\n" in value:
            lines.append(f'{key}: """')
            lines.append(value)
            lines.append('"""')
        elif isinstance(value, str):
            if value.startswith(("-", "[", "{")) or ":" in value:
                lines.append(f'{key}: "{value}"')
            else:
                lines.append(f"{key}: {value}")
        else:
            lines.append(f"{key}: {value}")
    lines.append("---")
    lines.append("")
    if body:
        lines.append(body)
    return "\n".join(lines)


def _export_json(fields: Dict[str, Any]) -> str:
    """Export entity to JSON format."""
    cleaned = {k: v for k, v in fields.items() if not k.startswith("_")}
    return json.dumps(cleaned, indent=2)


def _export_csv(fields: Dict[str, Any]) -> str:
    """Export entity to CSV format."""
    output = StringIO()
    writer = csv.writer(output)

    # Header
    headers = [k for k in fields.keys() if not k.startswith("_")]
    writer.writerow(headers)

    # Row
    row = [fields.get(h, "") for h in headers]
    writer.writerow(row)

    return output.getvalue()


def _export_html(fields: Dict[str, Any], body: str = "") -> str:
    """Export entity to HTML format."""
    template = """<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{title}</title>
    <style>
        body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; }}
        .entity-header {{ border-bottom: 2px solid #333; padding-bottom: 10px; margin-bottom: 20px; }}
        .field {{ margin-bottom: 15px; }}
        .field-name {{ font-weight: bold; color: #666; }}
        .field-value {{ margin-left: 10px; }}
        pre {{ background: #f4f4f4; padding: 10px; border-radius: 4px; overflow-x: auto; }}
    </style>
</head>
<body>
    <h1 class="entity-header">{{name}}</h1>
    {fields}
    {body}
</body>
</html>"""

    fields_html = ""
    for key, value in fields.items():
        if key.startswith("_"):
            continue
        if isinstance(value, list):
            value = ", ".join(str(v) for v in value)
        elif isinstance(value, dict):
            value = json.dumps(value, indent=2)

        fields_html += f"""
    <div class="field">
        <span class="field-name">{key}:</span>
        <span class="field-value">{value}</span>
    </div>
"""

    return template.replace("{title}", fields.get("name", "Entity")).replace("{name}", fields.get("name", "Entity")).replace("{fields}", fields_html).replace("{body}", f"<div class=\"field\"><span class=\"field-name\">Body:</span><span class=\"field-value\"><pre>{body}</pre></span></div>" if body else "")


def _export_pdf(fields: Dict[str, Any], body: str = "", template_path: Optional[str] = None) -> bytes:
    """Export entity to PDF format (placeholder - requires Playwright)."""
    import pdfkit

    html = _export_html(fields, body)

    # Load custom template if provided
    if template_path and os.path.exists(template_path):
        with open(template_path, 'r') as f:
            html = f.read().replace("{{name}}", fields.get("name", "Entity"))

    # Convert to PDF
    pdf_data = pdfkit.from_string(html, False)

    return pdf_data


def _export_docx(fields: Dict[str, Any], body: str = "", template_path: Optional[str] = None) -> bytes:
    """Export entity to DOCX format."""
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Apply template if provided
    if template_path and os.path.exists(template_path):
        doc = Document(template_path)
    else:
        # Default styling
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Segoe UI'
        font.size = Pt(11))

    # Add title
    title = doc.add_heading(fields.get("name", "Entity"), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add fields
    for key, value in fields.items():
        if key.startswith("_"):
            continue

        doc.add_paragraph(key, style='Heading 2')
        if isinstance(value, list):
            value = ", ".join(str(v) for v in value)
        elif isinstance(value, dict):
            value = json.dumps(value, indent=2)

        doc.add_paragraph(value)

    # Add body
    if body:
        doc.add_paragraph("Body", style='Heading 2')
        doc.add_paragraph(body)

    # Save to bytes
    from io import BytesIO
    output = BytesIO()
    doc.save(output)
    return output.getvalue()


def _export_verse(fields: Dict[str, Any], body: str = "") -> str:
    """Export entity to Verse format."""
    lines = [
        "# Auto-generated by StudioBrain Import/Export Pipeline",
        f"# Entity ID: {fields.get('id', 'unnamed')}",
        f"# Generated: {datetime.utcnow().isoformat()}Z",
        "#",
        f"using {{ /Verse.org/Simulation }}",
        f"",
    ]

    # Generate class from fields
    class_name = _sanitize_entity_id(fields.get("id", "unnamed")).capitalize() + "Data"
    lines.append(f"{class_name} := class:")
    lines.append(f'    Name : string = "{fields.get("name", "")}"')

    for key, value in fields.items():
        if key.startswith("_"):
            continue

        vtype = "string"
        if isinstance(value, bool):
            vtype = "logic"
        elif isinstance(value, int):
            vtype = "int"
        elif isinstance(value, float):
            vtype = "float"
        elif isinstance(value, list):
            vtype = "[]string"

        safe_key = re.sub(r"[^a-zA-Z0-9_]", "_", key.lower().strip())
        default = {"logic": "false", "int": "0", "float": "0.0"}.get(vtype, '""')

        lines.append(f"    {safe_key} : {vtype} = {default}")

    lines.append("")

    # Generate instance
    var_name = _sanitize_entity_id(fields.get("id", "unnamed"))
    lines.append(f"{var_name} := {class_name}:")

    for key, value in fields.items():
        if key.startswith("_"):
            continue

        safe_key = re.sub(r"[^a-zA-Z0-9_]", "_", key.lower().strip())
        vtype = "string"

        if isinstance(value, bool):
            vval = "true" if value else "false"
        elif isinstance(value, int):
            vval = str(value)
        elif isinstance(value, float):
            vval = str(value)
        elif isinstance(value, list):
            items = ", ".join(f'"{str(v)}"' for v in value)
            vval = f"array{{{items}}}"
        elif isinstance(value, str):
            escaped = value.replace("\\", "\\\\").replace('"', '\\"')
            vval = f'"{escaped}"'
        else:
            vval = '""'

        lines.append(f"    {safe_key} = {vval}")

    return "\n".join(lines)


# ─── Request Models ───────────────────────────────────────────────────


class ImportRequest(BaseModel):
    """Request body for import operations."""
    format: Optional[str] = None
    entity_type: Optional[str] = None
    entity_id: Optional[str] = None
    overwrite: bool = False
    content: Optional[str] = None
    content_base64: Optional[str] = None


class ImportBatchRequest(BaseModel):
    """Request body for batch import."""
    entities: List[ImportRequest]
    batch_size: Optional[int] = None


class ExportRequest(BaseModel):
    """Request body for export operations."""
    entity_type: str
    entity_id: str
    format: Optional[str] = None
    include_body: bool = True
    template_path: Optional[str] = None


class ExportBatchRequest(BaseModel):
    """Request body for batch export."""
    entity_type: Optional[str] = None
    entity_ids: List[str] = []
    format: Optional[str] = None


class ConflictResolveRequest(BaseModel):
    """Request for AI conflict resolution."""
    existing: Dict[str, Any]
    incoming: Dict[str, Any]
    entity_type: str


class ImportProfile(BaseModel):
    """Import profile configuration."""
    name: str
    format: Optional[str] = None
    entity_type: Optional[str] = None
    overwrite: bool = False
    ai_merge: bool = True


class ExportProfile(BaseModel):
    """Export profile configuration."""
    name: str
    format: Optional[str] = None
    include_images: bool = True
    include_relationships: bool = True
    template_path: Optional[str] = None


# ─── API Routes ───────────────────────────────────────────────────────


@router.get("/")
async def plugin_status():
    """Plugin status and configuration overview."""
    from services.plugin_settings_service import get_all_settings

    settings = get_all_settings("import-export-pipeline")

    # Count entities by type
    brains_root = _get_brains_root()
    entity_counts = {}
    for etype, dir_name in ENTITY_TYPE_DIRS.items():
        entity_dir = brains_root / dir_name
        if entity_dir.is_dir():
            entity_counts[etype] = sum(1 for d in entity_dir.iterdir() if d.is_dir())

    return {
        "plugin": "import-export-pipeline",
        "version": "1.0.0",
        "status": "ok",
        "settings": settings,
        "available_entity_types": list(ENTITY_TYPE_DIRS.keys()),
        "entity_counts": entity_counts,
        "import_formats": list(IMPORT_FORMATS.keys()),
        "export_formats": list(EXPORT_FORMATS.keys()),
    }


@router.get("/entity-types")
async def list_entity_types():
    """List all available entity types and their counts."""
    brains_root = _get_brains_root()
    result = []
    for etype, dir_name in ENTITY_TYPE_DIRS.items():
        entity_dir = brains_root / dir_name
        entity_ids = []
        if entity_dir.is_dir():
            entity_ids = [d.name for d in sorted(entity_dir.iterdir()) if d.is_dir()]

        result.append({
            "entity_type": etype,
            "directory": dir_name,
            "count": len(entity_ids),
            "entity_ids": entity_ids[:50],  # Cap at 50 for response size
        })
    return {"entity_types": result}


@router.get("/entities/{entity_type}")
async def list_entities_of_type(entity_type: str):
    """List all entities of a specific type."""
    if entity_type not in ENTITY_TYPE_DIRS:
        raise HTTPException(status_code=400, detail=f"Unknown entity type: {entity_type}")

    brains_root = _get_brains_root()
    entity_dir = brains_root / ENTITY_TYPE_DIRS[entity_type]

    if not entity_dir.is_dir():
        return {
            "entity_type": entity_type,
            "count": 0,
            "entity_ids": [],
        }

    entity_ids = [d.name for d in sorted(entity_dir.iterdir()) if d.is_dir()]
    return {
        "entity_type": entity_type,
        "count": len(entity_ids),
        "entity_ids": entity_ids,
    }


@router.post("/import")
async def import_entity(req: ImportRequest, background_tasks: BackgroundTasks):
    """Import a single entity from various formats."""
    # Get content
    if req.content:
        content = req.content
    elif req.content_base64:
        import base64
        content = base64.b64decode(req.content_base64).decode('utf-8')
    else:
        raise HTTPException(status_code=400, detail="No content provided")

    # Detect format if not specified
    fmt = req.format or _detect_format_from_content(content)
    if req.entity_id:
        fmt = req.format or _detect_format_from_filename(req.entity_id) or fmt

    # Import based on format
    try:
        if fmt == "markdown":
            result = await _import_markdown(content, req.entity_type)
        elif fmt == "json":
            result = await _import_json(content, req.entity_type)
        elif fmt == "csv":
            result = await _import_csv(content, req.entity_type)
        elif fmt == "opml":
            result = await _import_opml(content, req.entity_type)
        elif fmt == "verse":
            result = await _import_verse(content, req.entity_type)
        elif fmt == "zip":
            # For ZIP, we need bytes
            result = await _import_zip(content.encode('utf-8') if isinstance(content, str) else content, req.entity_type)
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported import format: {fmt}")

        # Process the result
        if isinstance(result, list):
            # Multiple entities from CSV or ZIP
            tasks = []
            for r in result:
                task = _import_entity(r, req.entity_type, req.entity_id, req.overwrite)
                tasks.append(task)
            results = await asyncio.gather(*tasks)
            return {
                "imported": len(results),
                "results": results,
            }
        else:
            # Single entity
            result = await _import_entity(result, req.entity_type, req.entity_id, req.overwrite)
            return result

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Import failed: {e}")
        raise HTTPException(status_code=500, detail=f"Import failed: {str(e)}")


@router.post("/import/batch")
async def batch_import(req: ImportBatchRequest):
    """Import multiple entities in batch."""
    batch_size = req.batch_size or _get_setting("batch_size", 10)
    results = []
    errors = []

    # Process in batches
    for i in range(0, len(req.entities), batch_size):
        batch = req.entities[i:i + batch_size]
        tasks = []

        for entity_req in batch:
            # Get content
            if entity_req.content:
                content = entity_req.content
            elif entity_req.content_base64:
                import base64
                content = base64.b64decode(entity_req.content_base64).decode('utf-8')
            else:
                errors.append({"index": len(results), "error": "No content provided"})
                continue

            # Import based on format
            try:
                fmt = entity_req.format or _detect_format_from_content(content)

                if fmt == "markdown":
                    result = await _import_markdown(content, entity_req.entity_type)
                elif fmt == "json":
                    result = await _import_json(content, entity_req.entity_type)
                elif fmt == "verse":
                    result = await _import_verse(content, entity_req.entity_type)
                else:
                    errors.append({"index": len(results), "error": f"Unsupported format: {fmt}"})
                    continue

                if isinstance(result, list):
                    for r in result:
                        task = _import_entity(r, entity_req.entity_type, entity_req.entity_id, entity_req.overwrite)
                        tasks.append(task)
                else:
                    task = _import_entity(result, entity_req.entity_type, entity_req.entity_id, entity_req.overwrite)
                    tasks.append(task)

            except Exception as e:
                errors.append({"index": len(results), "error": str(e)})

        batch_results = await asyncio.gather(*tasks, return_exceptions=True)
        for j, r in enumerate(batch_results):
            if isinstance(r, Exception):
                errors.append({"index": len(results) + j, "error": str(r)})
            else:
                results.append(r)

    return {
        "total_imported": len(results),
        "total_errors": len(errors),
        "results": results,
        "errors": errors,
    }


@router.post("/import/file")
async def import_file(
    file: str = Query(..., description="Base64-encoded file content"),
    filename: str = Query(..., description="Original filename for format detection"),
    entity_type: Optional[str] = Query(None),
    entity_id: Optional[str] = Query(None),
    overwrite: bool = Query(False),
):
    """Import an entity from an uploaded file."""
    import base64

    try:
        content = base64.b64decode(file).decode('utf-8')
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to decode file: {e}")

    # Detect format from filename
    detected_format = _detect_format_from_filename(filename)

    if detected_format == "zip":
        result = await _import_zip(content.encode('utf-8') if isinstance(content, str) else content, entity_type)
    elif detected_format == "markdown":
        result = await _import_markdown(content, entity_type)
    elif detected_format == "json":
        result = await _import_json(content, entity_type)
    elif detected_format == "csv":
        result = await _import_csv(content, entity_type)
    elif detected_format == "verse":
        result = await _import_verse(content, entity_type)
    else:
        # Try auto-detection
        detected_format = _detect_format_from_content(content)
        if detected_format == "markdown":
            result = await _import_markdown(content, entity_type)
        elif detected_format == "json":
            result = await _import_json(content, entity_type)
        else:
            result = await _import_markdown(content, entity_type)

    if isinstance(result, list):
        tasks = []
        for r in result:
            task = _import_entity(r, entity_type, entity_id, overwrite)
            tasks.append(task)
        results = await asyncio.gather(*tasks)
        return {
            "imported": len(results),
            "results": results,
        }
    else:
        result = await _import_entity(result, entity_type, entity_id, overwrite)
        return result


@router.post("/export")
async def export_entity(req: ExportRequest):
    """Export a single entity to the specified format."""
    entity = _read_entity(req.entity_type, req.entity_id)
    if not entity:
        raise HTTPException(status_code=404, detail=f"Entity not found: {req.entity_type}/{req.entity_id}")

    fmt = req.format or _get_setting("default_export_format", "markdown")
    body = entity.get("_body", "") if req.include_body else ""

    # Generate export based on format
    if fmt == "markdown":
        content = _export_markdown(entity, body)
        return {
            "format": "markdown",
            "content": content,
            "content_type": "text/markdown",
        }
    elif fmt == "json":
        content = _export_json(entity)
        return {
            "format": "json",
            "content": content,
            "content_type": "application/json",
        }
    elif fmt == "csv":
        content = _export_csv(entity)
        return {
            "format": "csv",
            "content": content,
            "content_type": "text/csv",
        }
    elif fmt == "html":
        content = _export_html(entity, body)
        return {
            "format": "html",
            "content": content,
            "content_type": "text/html",
        }
    elif fmt == "pdf":
        content = _export_pdf(entity, body, req.template_path)
        return {
            "format": "pdf",
            "content": base64.b64encode(content).decode('utf-8'),
            "content_type": "application/pdf",
        }
    elif fmt == "docx":
        content = _export_docx(entity, body, req.template_path)
        return {
            "format": "docx",
            "content": base64.b64encode(content).decode('utf-8'),
            "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        }
    elif fmt == "verse":
        content = _export_verse(entity, body)
        return {
            "format": "verse",
            "content": content,
            "content_type": "text/plain",
        }
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported export format: {fmt}")


@router.post("/export/batch")
async def batch_export(req: ExportBatchRequest):
    """Export multiple entities in batch."""
    if not req.entity_type and not req.entity_ids:
        return {
            "exported": 0,
            "results": [],
            "errors": ["No entities specified"],
        }

    # Determine entity type and IDs
    if req.entity_type:
        # All entities of this type
        if req.entity_ids:
            entity_ids = req.entity_ids
        else:
            entity_ids = [d.name for d in (_get_brains_root() / ENTITY_TYPE_DIRS.get(req.entity_type, "")).iterdir() if d.is_dir()]
    else:
        entity_ids = req.entity_ids

    fmt = req.format or _get_setting("default_export_format", "markdown")
    results = []
    errors = []

    for eid in entity_ids:
        try:
            entity = _read_entity(req.entity_type or "", eid)
            if not entity:
                errors.append({"entity_id": eid, "error": "Entity not found"})
                continue

            body = entity.get("_body", "")

            if fmt == "markdown":
                content = _export_markdown(entity, body)
            elif fmt == "json":
                content = _export_json(entity)
            elif fmt == "csv":
                content = _export_csv(entity)
            elif fmt == "html":
                content = _export_html(entity, body)
            elif fmt == "verse":
                content = _export_verse(entity, body)
            else:
                errors.append({"entity_id": eid, "error": f"Unsupported format: {fmt}"})
                continue

            results.append({
                "entity_type": entity.get("_entity_type", req.entity_type),
                "entity_id": eid,
                "format": fmt,
                "content": content,
            })

        except Exception as e:
            errors.append({"entity_id": eid, "error": str(e)})

    return {
        "exported": len(results),
        "total_attempted": len(entity_ids),
        "errors": len(errors),
        "results": results,
        "errors_list": errors,
    }


@router.post("/export/profiles")
async def list_export_profiles():
    """List available export profiles."""
    return {
        "profiles": [
            {
                "name": "default",
                "format": "markdown",
                "include_images": True,
                "include_relationships": True,
            },
            {
                "name": "pdf",
                "format": "pdf",
                "include_images": True,
                "include_relationships": True,
            },
            {
                "name": "docx",
                "format": "docx",
                "include_images": True,
                "include_relationships": True,
            },
            {
                "name": "json",
                "format": "json",
                "include_images": False,
                "include_relationships": False,
            },
            {
                "name": "verse",
                "format": "verse",
                "include_images": False,
                "include_relationships": False,
            },
        ]
    }


@router.post("/export/profiles/{profile_name}")
async def create_export_profile(profile_name: str, profile: ExportProfile):
    """Create or update an export profile."""
    return {
        "profile": {
            "name": profile_name,
            "format": profile.format,
            "include_images": profile.include_images,
            "include_relationships": profile.include_relationships,
            "template_path": profile.template_path,
        }
    }


@router.post("/import/profiles")
async def list_import_profiles():
    """List available import profiles."""
    return {
        "profiles": [
            {
                "name": "default",
                "format": "auto",
                "overwrite": False,
                "ai_merge": True,
            },
            {
                "name": "force-merge",
                "format": "auto",
                "overwrite": False,
                "ai_merge": True,
            },
            {
                "name": "overwrite",
                "format": "auto",
                "overwrite": True,
                "ai_merge": False,
            },
        ]
    }


@router.post("/resolve-conflict")
async def resolve_conflict(req: ConflictResolveRequest):
    """Resolve a merge conflict using AI or simple strategy."""
    result = await _ai_resolve_conflict(req.existing, req.incoming, req.entity_type)
    return {"resolved": result}


@router.get("/audit")
async def get_import_export_audit(
    start_date: Optional[str] = Query(None),
    end_date: Optional[str] = Query(None),
    entity_type: Optional[str] = Query(None),
):
    """Get audit log of import/export operations."""
    # In production, this would query a database
    return {
        "audit_log": [],
        "summary": {
            "total_imports": 0,
            "total_exports": 0,
            "total_entities_processed": 0,
        },
    }


@router.get("/formats")
async def list_formats():
    """List all supported import and export formats."""
    return {
        "import": {
            "markdown": {"extensions": [".md", ".markdown"], "description": "Markdown with YAML frontmatter"},
            "json": {"extensions": [".json"], "description": "JSON data"},
            "csv": {"extensions": [".csv"], "description": "Comma-separated values"},
            "opml": {"extensions": [".opml"], "description": "Outline processor format"},
            "verse": {"extensions": [".verse", ".vrs"], "description": "Verse (UEFN) code"},
            "zip": {"extensions": [".zip"], "description": "ZIP archive of multiple files"},
        },
        "export": {
            "markdown": {"extensions": [".md"], "description": "Markdown with YAML frontmatter"},
            "json": {"extensions": [".json"], "description": "JSON data"},
            "csv": {"extensions": [".csv"], "description": "Comma-separated values"},
            "html": {"extensions": [".html"], "description": "HTML page"},
            "pdf": {"extensions": [".pdf"], "description": "PDF document"},
            "docx": {"extensions": [".docx"], "description": "Microsoft Word document"},
            "verse": {"extensions": [".verse"], "description": "Verse (UEFN) code"},
        },
    }
